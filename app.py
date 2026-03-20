"""
app.py
======
Streamlit app for 4-camera oblique aerial survey flight planning.

Run:
    streamlit run app.py
"""

import json
import math
import os
from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import streamlit as st

from geometry import (
    normalize_tilt_angle,
    calculate_camera_solution,
    calculate_multicamera_solution,
    half_fov_deg,
    pixel_size_mm,
    m_to_unit,
    unit_to_m,
    mm_to_unit,
)

# ---------------------------------------------------------------------------
# Page config
# ---------------------------------------------------------------------------

st.set_page_config(
    page_title="Oblique Survey Planner",
    page_icon="✈️",
    layout="wide",
)

# ---------------------------------------------------------------------------
# Presets
# ---------------------------------------------------------------------------

DEFAULT_PRESETS = {
    "Sony A7R IV": {
        "sensor_width_mm": 35.7,
        "sensor_height_mm": 23.8,
        "image_width_px": 9504,
        "image_height_px": 6336,
        "focal_length_mm": 35.0,
    },
    "Sony A7R V": {
        "sensor_width_mm": 35.7,
        "sensor_height_mm": 23.8,
        "image_width_px": 9504,
        "image_height_px": 6336,
        "focal_length_mm": 35.0,
    },
    "Phase One iXM-100": {
        "sensor_width_mm": 53.4,
        "sensor_height_mm": 40.0,
        "image_width_px": 11664,
        "image_height_px": 8750,
        "focal_length_mm": 50.0,
    },
}

PRESET_FILE = Path("presets.json")

# ---------------------------------------------------------------------------
# Overlap presets
# ---------------------------------------------------------------------------

OVERLAP_PRESETS = {
    "(custom)": None,
    "Oblique browse – efficient": {"forward": 25, "sidelap": 5},
    "Oblique browse – standard": {"forward": 30, "sidelap": 10},
    "Oblique browse – urban safe": {"forward": 35, "sidelap": 15},
    "Photogrammetry": {"forward": 60, "sidelap": 35},
    "High-rise / lean control": {"forward": 80, "sidelap": 80},
}


def load_presets() -> dict:
    if PRESET_FILE.exists():
        try:
            with open(PRESET_FILE) as f:
                saved = json.load(f)
            return {**DEFAULT_PRESETS, **saved}
        except Exception:
            pass
    return dict(DEFAULT_PRESETS)


def save_preset(name: str, data: dict):
    existing = {}
    if PRESET_FILE.exists():
        try:
            with open(PRESET_FILE) as f:
                existing = json.load(f)
        except Exception:
            pass
    existing[name] = data
    with open(PRESET_FILE, "w") as f:
        json.dump(existing, f, indent=2)


def save_scenario(scenario: dict, filename: str = "scenario.json"):
    with open(filename, "w") as f:
        json.dump(scenario, f, indent=2)
    return filename


# ---------------------------------------------------------------------------
# Unit helpers for display
# ---------------------------------------------------------------------------

DIST_UNITS = ["m", "ft", "cm"]
SENSOR_UNITS = ["mm", "cm"]

def fmt(val, unit="m", decimals=2):
    return f"{m_to_unit(val, unit):.{decimals}f} {unit}"

def fmt_gsd(val_m, unit="cm", decimals=1):
    return f"{m_to_unit(val_m, unit):.{decimals}f} {unit}/px"

def fmt_mm(val_mm, unit="mm", decimals=3):
    return f"{mm_to_unit(val_mm, unit):.{decimals}f} {unit}"


# ---------------------------------------------------------------------------
# Sidebar inputs
# ---------------------------------------------------------------------------

st.sidebar.title("✈️ Oblique Survey Planner")
st.sidebar.markdown("---")

# --- Presets ---
presets = load_presets()
st.sidebar.subheader("📷 Camera Preset")
preset_names = list(presets.keys())
selected_preset = st.sidebar.selectbox("Load preset", ["(custom)"] + preset_names)

if selected_preset != "(custom)":
    preset = presets[selected_preset]
else:
    preset = DEFAULT_PRESETS["Sony A7R IV"]

# --- Camera Setup ---
st.sidebar.subheader("Camera Setup")
camera_name = st.sidebar.text_input("Camera name", value=selected_preset if selected_preset != "(custom)" else "My Camera")

col1, col2 = st.sidebar.columns(2)
with col1:
    sensor_w_mm = st.sidebar.number_input("Sensor width (mm)", value=float(preset["sensor_width_mm"]), min_value=1.0, max_value=200.0, step=0.1)
    sensor_h_mm = st.sidebar.number_input("Sensor height (mm)", value=float(preset["sensor_height_mm"]), min_value=1.0, max_value=200.0, step=0.1)
with col2:
    img_w_px = st.sidebar.number_input("Image width (px)", value=int(preset["image_width_px"]), min_value=100, max_value=100000, step=100)
    img_h_px = st.sidebar.number_input("Image height (px)", value=int(preset["image_height_px"]), min_value=100, max_value=100000, step=100)

focal_mm = st.sidebar.number_input("Focal length (mm)", value=float(preset["focal_length_mm"]), min_value=1.0, max_value=2000.0, step=1.0)

arrangement = st.sidebar.selectbox(
    "Camera arrangement",
    ["4 oblique", "2 oblique", "Single nadir"],
    index=0,
)
arr_map = {"4 oblique": "4_oblique", "2 oblique": "2_oblique", "Single nadir": "single_nadir"}
arrangement_key = arr_map[arrangement]

angle_convention = st.sidebar.selectbox(
    "Tilt angle measured from",
    ["Nadir (0° = straight down)", "Horizontal (0° = level)"],
)
convention_key = "nadir" if "Nadir" in angle_convention else "horiz"

if arrangement_key == "single_nadir":
    tilt_input = 0.0
    st.sidebar.info("Single nadir: tilt fixed at 0°")
else:
    label = "Tilt angle (°) — from nadir" if convention_key == "nadir" else "Tilt angle (°) — from horizontal"
    tilt_input = st.sidebar.slider(label, min_value=0.0, max_value=85.0, value=35.0, step=0.5)

tilt_from_nadir = normalize_tilt_angle(tilt_input, convention_key)

st.sidebar.markdown("---")

# --- Flight Setup ---
st.sidebar.subheader("Flight Setup")
dist_unit = st.sidebar.selectbox("Distance display unit", ["m", "ft", "cm"], index=0)

altitude_input = st.sidebar.number_input(
    f"Altitude AGL ({dist_unit})",
    value=m_to_unit(1000.0, dist_unit),
    min_value=1.0,
    max_value=m_to_unit(10000.0, dist_unit),
    step=m_to_unit(10.0, dist_unit),
)
altitude_m = unit_to_m(altitude_input, dist_unit)

speed_ms = st.sidebar.number_input("Aircraft speed (m/s)", value=50.0, min_value=1.0, max_value=200.0, step=1.0)
speed_kts = speed_ms * 1.94384
st.sidebar.caption(f"≈ {speed_kts:.1f} knots")

reciprocal_flying = st.sidebar.checkbox("Reciprocal flying (bidirectional strips)", value=True)

st.sidebar.markdown("---")

# --- Overlap Settings ---
st.sidebar.subheader("Overlap Settings")

# Overlap preset dropdown — determines the default values fed into the sliders below.
# Selecting a named preset populates both sliders; the user can still drag them freely afterward.
selected_overlap_preset = st.sidebar.selectbox(
    "Overlap preset",
    list(OVERLAP_PRESETS.keys()),
    index=0,
    key="overlap_preset_select",
    help="Choose a named overlap configuration to auto-fill the sliders, then adjust as needed.",
)

# Seed slider session_state on very first run
if "forward_overlap_slider" not in st.session_state:
    st.session_state["forward_overlap_slider"] = 60
if "sidelap_slider" not in st.session_state:
    st.session_state["sidelap_slider"] = 30

# Detect when the user picks a different preset and push the new values into
# session_state BEFORE the sliders render. We track the previously-applied
# preset in its own session_state key so we only update on an actual change.
if st.session_state.get("_last_overlap_preset") != selected_overlap_preset:
    st.session_state["_last_overlap_preset"] = selected_overlap_preset
    overlap_preset_values = OVERLAP_PRESETS[selected_overlap_preset]
    if overlap_preset_values is not None:
        st.session_state["forward_overlap_slider"] = overlap_preset_values["forward"]
        st.session_state["sidelap_slider"] = overlap_preset_values["sidelap"]

# Sliders read from session_state[key]; no value= argument needed or wanted.
forward_overlap_pct = st.sidebar.slider(
    "Forward overlap (%)", 10, 95,
    key="forward_overlap_slider",
)
sidelap_pct = st.sidebar.slider(
    "Sidelap (%)", 10, 95,
    key="sidelap_slider",
)
forward_overlap = forward_overlap_pct / 100.0
sidelap = sidelap_pct / 100.0

st.sidebar.markdown("---")

# --- Save / Load ---
st.sidebar.subheader("💾 Save Scenario")
scenario_name = st.sidebar.text_input("Scenario name", value="my_survey")
if st.sidebar.button("Save scenario as JSON"):
    scenario = {
        "camera_name": camera_name,
        "sensor_width_mm": sensor_w_mm,
        "sensor_height_mm": sensor_h_mm,
        "image_width_px": img_w_px,
        "image_height_px": img_h_px,
        "focal_length_mm": focal_mm,
        "arrangement": arrangement,
        "tilt_input_deg": tilt_input,
        "angle_convention": convention_key,
        "altitude_m": altitude_m,
        "speed_ms": speed_ms,
        "forward_overlap_pct": forward_overlap_pct,
        "sidelap_pct": sidelap_pct,
        "reciprocal_flying": reciprocal_flying,
    }
    filename = f"{scenario_name}.json"
    save_scenario(scenario, filename)
    st.sidebar.success(f"Saved to {filename}")

st.sidebar.subheader("📥 Save Camera Preset")
new_preset_name = st.sidebar.text_input("Preset name", value=camera_name)
if st.sidebar.button("Save as preset"):
    save_preset(new_preset_name, {
        "sensor_width_mm": sensor_w_mm,
        "sensor_height_mm": sensor_h_mm,
        "image_width_px": img_w_px,
        "image_height_px": img_h_px,
        "focal_length_mm": focal_mm,
    })
    st.sidebar.success(f"Preset '{new_preset_name}' saved. Reload to see it.")


# ---------------------------------------------------------------------------
# Compute
# ---------------------------------------------------------------------------

# Determine camera layout
# For 4-oblique: cameras at +tilt (right) and -tilt (left), ×2 (forward/rear)
# For 2-oblique: one at +tilt, one at -tilt
# For single nadir: one camera at 0°

if arrangement_key == "4_oblique":
    tilts = [tilt_from_nadir, tilt_from_nadir, -tilt_from_nadir, -tilt_from_nadir]
    labels = ["Right-Fwd", "Right-Rear", "Left-Fwd", "Left-Rear"]
    # VERIFY: This assumes all 4 cameras point outward symmetrically.
    # In real 4-oblique rigs, cameras may be angled both across and along track.
    # Here we model all 4 as purely across-track oblique (simplified).
elif arrangement_key == "2_oblique":
    tilts = [tilt_from_nadir, -tilt_from_nadir]
    labels = ["Right", "Left"]
else:
    tilts = [0.0]
    labels = ["Nadir"]

cam_solutions = []
errors = []
for i, tilt in enumerate(tilts):
    try:
        sol = calculate_camera_solution(
            altitude_m=altitude_m,
            tilt_from_nadir_deg=abs(tilt),  # geometry is symmetric; sign = left/right
            sensor_width_mm=sensor_w_mm,
            sensor_height_mm=sensor_h_mm,
            image_width_px=img_w_px,
            image_height_px=img_h_px,
            focal_length_mm=focal_mm,
        )
        cam_solutions.append(sol)
    except Exception as e:
        errors.append(f"Camera {labels[i]}: {e}")

mc = None
if cam_solutions and not errors:
    try:
        mc = calculate_multicamera_solution(
            camera_solutions=cam_solutions,
            arrangement=arrangement_key,
            altitude_m=altitude_m,
            aircraft_speed_ms=speed_ms,
            forward_overlap_fraction=forward_overlap,
            sidelap_fraction=sidelap,
            reciprocal_flying=reciprocal_flying,
        )
    except Exception as e:
        errors.append(f"Multi-camera: {e}")


# ---------------------------------------------------------------------------
# Main page
# ---------------------------------------------------------------------------

st.title("✈️ Oblique Aerial Survey Planner")
st.caption(
    f"**Camera:** {camera_name}  |  **Arrangement:** {arrangement}  |  "
    f"**Tilt:** {tilt_from_nadir:.1f}° from nadir  |  "
    f"**Altitude:** {fmt(altitude_m, dist_unit)} AGL  |  "
    f"**Speed:** {speed_ms:.1f} m/s"
)

# --- Errors / warnings ---
if errors:
    for e in errors:
        st.error(e)

if mc and mc.warnings:
    with st.expander("⚠️ Warnings", expanded=True):
        for w in mc.warnings:
            st.warning(w)

if not cam_solutions:
    st.stop()

# ---------------------------------------------------------------------------
# Summary cards
# ---------------------------------------------------------------------------

st.subheader("System Summary")

if mc:
    c1, c2, c3, c4, c5 = st.columns(5)
    with c1:
        st.metric("Combined Swath", fmt(mc.combined_swath_m, dist_unit))
    with c2:
        st.metric("Line Spacing", fmt(mc.recommended_line_spacing_m, dist_unit))
    with c3:
        st.metric("Photo Spacing", fmt(mc.recommended_photo_spacing_m, dist_unit))
    with c4:
        st.metric("Exposure Interval", f"{mc.photo_interval_s:.2f} s")
    with c5:
        st.metric("Sidelap Achieved", f"{mc.sidelap_achieved*100:.1f}%")

    c6, c7, c8, c9, c10 = st.columns(5)
    with c6:
        sol0 = cam_solutions[0]
        st.metric("Near-Edge GSD", fmt_gsd(sol0.near_gsd_m, "cm"))
    with c7:
        st.metric("Centre GSD", fmt_gsd(sol0.centre_gsd_m, "cm"))
    with c8:
        st.metric("Far-Edge GSD", fmt_gsd(sol0.far_gsd_m, "cm"))
    with c9:
        st.metric("Fwd Overlap (centre)", f"{mc.forward_overlap_centre*100:.1f}%")
    with c10:
        rec = "✅ Yes" if mc.reciprocal_recommended else "Not required"
        st.metric("Reciprocal Flying", rec)

st.markdown("---")

# ---------------------------------------------------------------------------
# Near-edge viewing angles
# ---------------------------------------------------------------------------

st.subheader("Frame Viewing Angles")
st.caption(
    "Camera ray angles to the near edge and centre of the frame, shown from two references: "
    "from vertical (nadir) and from horizontal. "
    "Values are for the first camera in the arrangement (all cameras share the same tilt)."
)

if cam_solutions:
    sol0 = cam_solutions[0]
    near_from_nadir  = sol0.near_angle_deg
    near_from_horiz  = 90.0 - near_from_nadir
    ctr_from_nadir   = sol0.centre_angle_deg
    ctr_from_horiz   = 90.0 - ctr_from_nadir

    # --- Near edge row ---
    st.markdown("**Near edge of frame**")
    va1, va2, va3 = st.columns(3)
    with va1:
        st.metric(
            "Near-Edge — from vertical (nadir)",
            f"{near_from_nadir:.1f}°",
            help="0° = camera pointing straight down. 90° = camera pointing at the horizon.",
        )
    with va2:
        st.metric(
            "Near-Edge — from horizontal",
            f"{near_from_horiz:.1f}°",
            help="90° = camera pointing straight down. 0° = camera pointing at the horizon.",
        )
    with va3:
        if near_from_nadir < 10:
            feel = "🟢 Near-vertical — minimal oblique distortion"
        elif near_from_nadir < 30:
            feel = "🟡 Mildly oblique"
        elif near_from_nadir < 50:
            feel = "🟠 Moderately oblique"
        elif near_from_nadir < 70:
            feel = "🔴 Highly oblique — significant GSD variation"
        else:
            feel = "🔴 Extreme oblique — near-grazing geometry"
        st.metric("Near-Edge Obliqueness", feel)

    # --- Centre row ---
    st.markdown("**Centre of frame**")
    vb1, vb2, vb3 = st.columns(3)
    with vb1:
        st.metric(
            "Centre — from vertical (nadir)",
            f"{ctr_from_nadir:.1f}°",
            help="0° = camera pointing straight down. 90° = camera pointing at the horizon.",
        )
    with vb2:
        st.metric(
            "Centre — from horizontal",
            f"{ctr_from_horiz:.1f}°",
            help="90° = camera pointing straight down. 0° = camera pointing at the horizon.",
        )
    with vb3:
        if ctr_from_nadir < 10:
            feel_ctr = "🟢 Near-vertical — minimal oblique distortion"
        elif ctr_from_nadir < 30:
            feel_ctr = "🟡 Mildly oblique"
        elif ctr_from_nadir < 50:
            feel_ctr = "🟠 Moderately oblique"
        elif ctr_from_nadir < 70:
            feel_ctr = "🔴 Highly oblique — significant GSD variation"
        else:
            feel_ctr = "🔴 Extreme oblique — near-grazing geometry"
        st.metric("Centre Obliqueness", feel_ctr)

st.markdown("---")

# ---------------------------------------------------------------------------
# Per-camera results table
# ---------------------------------------------------------------------------

st.subheader("Per-Camera Geometry")

table_data = []
for i, (lbl, sol) in enumerate(zip(labels, cam_solutions)):
    side = "Right" if tilts[i] >= 0 else "Left"
    table_data.append({
        "Camera": lbl,
        "Side": side,
        f"Near Edge ({dist_unit})": round(m_to_unit(sol.near_edge_m, dist_unit), 1),
        f"Centre ({dist_unit})": round(m_to_unit(sol.centre_m, dist_unit), 1),
        f"Far Edge ({dist_unit})": round(m_to_unit(sol.far_edge_m, dist_unit), 1),
        "Near Slant (m)": round(sol.near_slant_m, 1),
        "Centre Slant (m)": round(sol.centre_slant_m, 1),
        "Far Slant (m)": round(sol.far_slant_m, 1),
        "Near angle from vertical (°)": round(sol.near_angle_deg, 1),
        "Near angle from horizontal (°)": round(90.0 - sol.near_angle_deg, 1),
        "Centre angle from vertical (°)": round(sol.centre_angle_deg, 1),
        "Centre angle from horizontal (°)": round(90.0 - sol.centre_angle_deg, 1),
        "Near GSD (cm/px)": round(sol.near_gsd_m * 100, 2),
        "Centre GSD (cm/px)": round(sol.centre_gsd_m * 100, 2),
        "Far GSD (cm/px)": round(sol.far_gsd_m * 100, 2),
        f"Footprint Across ({dist_unit})": round(m_to_unit(sol.footprint_across_m, dist_unit), 1),
        f"Footprint Along ({dist_unit})": round(m_to_unit(sol.footprint_along_m, dist_unit), 1),
        "Pixel Size (µm)": round(sol.pixel_size_mm * 1000, 2),
        "Half FOV Across (°)": round(sol.half_fov_across_deg, 2),
    })

st.dataframe(table_data, use_container_width=True)

st.markdown("---")

# ---------------------------------------------------------------------------
# Cross-section diagram
# ---------------------------------------------------------------------------

st.subheader("Cross-Section View")
st.caption("Side view looking along the flight track. Shows camera rays to near edge, centre, and far edge on flat ground.")

fig_xs, ax_xs = plt.subplots(figsize=(12, 5))
fig_xs.patch.set_facecolor("#0e1117")
ax_xs.set_facecolor("#1a1d2e")

# Aircraft position
ac_x, ac_y = 0.0, altitude_m
ax_xs.scatter([ac_x], [ac_y], s=200, color="#f0c040", zorder=5, marker="^", label="Aircraft")
ax_xs.vlines(0, 0, altitude_m, color="#555", linewidth=1, linestyle="--", label="Nadir")

# Ground line
max_extent = max(abs(sol.far_edge_m) for sol in cam_solutions) * 1.15
ax_xs.axhline(0, color="#888", linewidth=1.5, label="Ground")
ax_xs.set_xlim(-max_extent, max_extent)
ax_xs.set_ylim(-0.05 * altitude_m, altitude_m * 1.1)

colors_right = ["#4fc3f7", "#29b6f6"]
colors_left = ["#ef9a9a", "#e57373"]

for i, (lbl, sol) in enumerate(zip(labels, cam_solutions)):
    is_right = tilts[i] >= 0
    col = colors_right[i % 2] if is_right else colors_left[i % 2]
    sign = 1.0 if is_right else -1.0

    near_x = sign * sol.near_edge_m
    ctr_x = sign * sol.centre_m
    far_x = sign * sol.far_edge_m

    # Draw rays
    for gx, style, label in [
        (near_x, ":", f"{lbl} near"),
        (ctr_x, "-", f"{lbl} centre"),
        (far_x, "--", f"{lbl} far"),
    ]:
        ax_xs.plot([ac_x, gx], [ac_y, 0], color=col, linestyle=style, linewidth=1.5, label=label)

    # Mark ground intercepts
    ax_xs.scatter([near_x, ctr_x, far_x], [0, 0, 0], color=col, s=50, zorder=4)

    # Annotate near/far
    ax_xs.annotate(f"{lbl}\nnear: {m_to_unit(sol.near_edge_m, dist_unit):.0f}{dist_unit}",
                   xy=(near_x, 0), xytext=(near_x, -0.03*altitude_m),
                   fontsize=7, color=col, ha="center", va="top")
    ax_xs.annotate(f"far: {m_to_unit(sol.far_edge_m, dist_unit):.0f}{dist_unit}",
                   xy=(far_x, 0), xytext=(far_x, -0.06*altitude_m),
                   fontsize=7, color=col, ha="center", va="top")

ax_xs.set_xlabel(f"Across-track distance ({dist_unit})", color="white")
ax_xs.set_ylabel(f"Altitude ({dist_unit})", color="white")
ax_xs.tick_params(colors="white")
ax_xs.spines[:].set_color("#555")
ax_xs.set_title("Cross-Section: Camera Rays to Ground", color="white", fontsize=12)

# Re-label axes in display unit
xticks = ax_xs.get_xticks()
ax_xs.set_xticklabels([f"{m_to_unit(t, dist_unit):.0f}" for t in xticks], color="white")
yticks = ax_xs.get_yticks()
ax_xs.set_yticklabels([f"{m_to_unit(t, dist_unit):.0f}" for t in yticks], color="white")

ax_xs.legend(loc="upper right", fontsize=7, framealpha=0.3, labelcolor="white", facecolor="#222")
fig_xs.tight_layout()
st.pyplot(fig_xs)

st.markdown("---")

# ---------------------------------------------------------------------------
# Plan-view diagram
# ---------------------------------------------------------------------------

st.subheader("Plan View — Strip Overlap")
st.caption(
    "Top-down view showing adjacent flight strips. "
    "Coloured bands show the ground footprint of each strip. "
    "Grey hatching shows the overlapping (shared) zone."
)

if mc:
    fig_pv, ax_pv = plt.subplots(figsize=(12, 6))
    fig_pv.patch.set_facecolor("#0e1117")
    ax_pv.set_facecolor("#1a1d2e")

    line_spacing = mc.recommended_line_spacing_m
    n_strips = 3  # show 3 adjacent strips

    # Use the first camera's near/far (right side) and mirror for left
    sol0 = cam_solutions[0]
    strip_colors = ["#1a3a5c", "#1a4a3c", "#3a1a5c"]
    edge_colors = ["#4fc3f7", "#4caf50", "#ab47bc"]

    for strip_idx in range(n_strips):
        track_y = strip_idx * line_spacing  # nadir track position in across-track axis

        # Footprint spans from near_edge to far_edge, centred on track_y
        # Right-side cameras
        right_near = track_y + sol0.near_edge_m
        right_far = track_y + sol0.far_edge_m

        along_half = sol0.footprint_along_m / 2.0
        rect = mpatches.FancyArrowPatch

        # Draw footprint rectangle in plan view
        # x-axis = along-track (flight direction), y-axis = across-track
        width = sol0.footprint_along_m
        height = sol0.far_edge_m - sol0.near_edge_m  # across-track

        rect_patch = plt.Rectangle(
            (-along_half, right_near),
            width, height,
            linewidth=1.5, edgecolor=edge_colors[strip_idx],
            facecolor=strip_colors[strip_idx], alpha=0.6,
            label=f"Strip {strip_idx+1} (track at {m_to_unit(track_y, dist_unit):.0f} {dist_unit})"
        )
        ax_pv.add_patch(rect_patch)

        # Left-side cameras (mirrored)
        if len(cam_solutions) > 1:
            left_near = track_y - sol0.far_edge_m
            left_far = track_y - sol0.near_edge_m
            rect_patch_l = plt.Rectangle(
                (-along_half, left_near),
                width, left_far - left_near,
                linewidth=1.5, edgecolor=edge_colors[strip_idx],
                facecolor=strip_colors[strip_idx], alpha=0.4,
            )
            ax_pv.add_patch(rect_patch_l)

        # Draw nadir track line
        ax_pv.axhline(track_y, color=edge_colors[strip_idx], linewidth=1, linestyle="-.", alpha=0.8)
        ax_pv.text(
            along_half * 1.02, track_y,
            f"Track {strip_idx+1}",
            color=edge_colors[strip_idx], fontsize=8, va="center"
        )

    # Show line spacing annotation
    ax_pv.annotate(
        "",
        xy=(0, line_spacing),
        xytext=(0, 0),
        arrowprops=dict(arrowstyle="<->", color="white", lw=1.5),
    )
    ax_pv.text(
        along_half * 0.05, line_spacing / 2,
        f"Line spacing\n{m_to_unit(line_spacing, dist_unit):.0f} {dist_unit}",
        color="white", fontsize=8, va="center",
        bbox=dict(facecolor="#333", alpha=0.6, pad=2)
    )

    ax_pv.set_xlim(-along_half * 1.3, along_half * 1.5)
    ax_pv.set_ylim(
        -sol0.far_edge_m * 0.5,
        (n_strips - 1) * line_spacing + sol0.far_edge_m * 1.3
    )
    ax_pv.set_xlabel(f"Along-track direction ({dist_unit})", color="white")
    ax_pv.set_ylabel(f"Across-track direction ({dist_unit})", color="white")
    ax_pv.tick_params(colors="white")
    ax_pv.spines[:].set_color("#555")
    ax_pv.set_title("Plan View: Adjacent Strips", color="white", fontsize=12)
    ax_pv.legend(loc="upper left", fontsize=8, framealpha=0.3, labelcolor="white", facecolor="#222")

    # Re-label in display units
    xticks_pv = ax_pv.get_xticks()
    ax_pv.set_xticklabels([f"{m_to_unit(t, dist_unit):.0f}" for t in xticks_pv], color="white")
    yticks_pv = ax_pv.get_yticks()
    ax_pv.set_yticklabels([f"{m_to_unit(t, dist_unit):.0f}" for t in yticks_pv], color="white")

    fig_pv.tight_layout()
    st.pyplot(fig_pv)

st.markdown("---")

# ---------------------------------------------------------------------------
# Overlap detail
# ---------------------------------------------------------------------------

if mc:
    st.subheader("Overlap Detail")
    oc1, oc2, oc3 = st.columns(3)
    with oc1:
        st.metric("Fwd Overlap — Near Edge", f"{mc.forward_overlap_near*100:.1f}%")
    with oc2:
        st.metric("Fwd Overlap — Centre", f"{mc.forward_overlap_centre*100:.1f}%")
    with oc3:
        st.metric("Fwd Overlap — Far Edge", f"{mc.forward_overlap_far*100:.1f}%")

st.markdown("---")

# ---------------------------------------------------------------------------
# Intermediate values / formula trace
# ---------------------------------------------------------------------------

st.subheader("Intermediate Values & Formula Trace")

with st.expander("Show intermediate calculations (first camera)", expanded=False):
    sol0 = cam_solutions[0]
    px_sz_mm = pixel_size_mm(sensor_w_mm, img_w_px)

    st.markdown("**Sensor & Pixel Geometry**")
    st.markdown(f"""
| Parameter | Formula | Value |
|-----------|---------|-------|
| Pixel size | `sensor_width / image_width` | `{sensor_w_mm} / {img_w_px}` = **{px_sz_mm*1000:.2f} µm** |
| Focal length (px) | `focal_mm / pixel_size_mm` | `{focal_mm} / {px_sz_mm:.4f}` = **{focal_mm/px_sz_mm:.1f} px** |
| Half FOV across | `atan(sensor_w / (2 × fl))` | `atan({sensor_w_mm} / {2*focal_mm:.1f})` = **{sol0.half_fov_across_deg:.2f}°** |
| Half FOV along | `atan(sensor_h / (2 × fl))` | `atan({sensor_h_mm} / {2*focal_mm:.1f})` = **{sol0.half_fov_along_deg:.2f}°** |
    """)

    st.markdown("**Ground Intercepts (flat terrain, pinhole model)**")
    st.markdown(f"""
| Ray | Angle from nadir | Formula | Ground distance | Slant range |
|-----|-----------------|---------|-----------------|-------------|
| Near edge | `{sol0.near_angle_deg:.2f}°` | `H × tan(θ - φ)` | **{m_to_unit(sol0.near_edge_m, dist_unit):.1f} {dist_unit}** | **{sol0.near_slant_m:.1f} m** |
| Centre    | `{sol0.centre_angle_deg:.2f}°` | `H × tan(θ)` | **{m_to_unit(sol0.centre_m, dist_unit):.1f} {dist_unit}** | **{sol0.centre_slant_m:.1f} m** |
| Far edge  | `{sol0.far_angle_deg:.2f}°` | `H × tan(θ + φ)` | **{m_to_unit(sol0.far_edge_m, dist_unit):.1f} {dist_unit}** | **{sol0.far_slant_m:.1f} m** |
    """)

    st.markdown("**GSD Calculation**")
    st.markdown(f"""
GSD formula (pinhole, flat ground):
```
GSD = (pixel_size_mm / focal_length_mm) × altitude / cos²(angle_from_nadir)
```

| Position | Angle | GSD |
|----------|-------|-----|
| Near edge | `{sol0.near_angle_deg:.2f}°` | **{sol0.near_gsd_m*100:.2f} cm/px** |
| Centre    | `{sol0.centre_angle_deg:.2f}°` | **{sol0.centre_gsd_m*100:.2f} cm/px** |
| Far edge  | `{sol0.far_angle_deg:.2f}°` | **{sol0.far_gsd_m*100:.2f} cm/px** |
    """)

    if mc:
        st.markdown("**Line & Photo Spacing**")
        st.markdown(f"""
| Parameter | Formula | Value |
|-----------|---------|-------|
| Combined swath | `far_edge_max - near_edge_min` | **{m_to_unit(mc.combined_swath_m, dist_unit):.1f} {dist_unit}** |
| Line spacing | `combined_swath × (1 - sidelap)` | `{m_to_unit(mc.combined_swath_m, dist_unit):.1f} × {1-sidelap:.2f}` = **{m_to_unit(mc.recommended_line_spacing_m, dist_unit):.1f} {dist_unit}** |
| Photo spacing | `footprint_along × (1 - fwd_overlap)` | `{m_to_unit(sol0.footprint_along_m, dist_unit):.1f} × {1-forward_overlap:.2f}` = **{m_to_unit(mc.recommended_photo_spacing_m, dist_unit):.1f} {dist_unit}** |
| Exposure interval | `photo_spacing / speed` | `{mc.recommended_photo_spacing_m:.1f} / {speed_ms:.1f}` = **{mc.photo_interval_s:.2f} s** |
        """)

st.markdown("---")

# ---------------------------------------------------------------------------
# Assumptions & Help
# ---------------------------------------------------------------------------

st.subheader("ℹ️ Assumptions & Help")

with st.expander("Click to expand assumptions and geometry notes"):
    st.markdown("""
### Camera Model
- **Pinhole camera model** — no lens distortion, no chromatic aberration.
- GSD is computed from sensor pixel size, focal length, and slant range.
- Pixel size = `sensor_dimension_mm / image_dimension_px` (assumes square pixels).

### Terrain Model
- **Flat terrain (v1)** — the ground is assumed to be a horizontal plane at the stated AGL altitude.
- Ground elevation variation is not yet accounted for in geometry calculations.
- For undulating terrain, use the altitude AGL above the *highest* ground point as a conservative estimate for GSD, and the altitude above the *lowest* point for conservative footprint sizing.

### Angle Convention
- **Tilt from nadir** — 0° = camera pointing straight down, 90° = camera pointing at horizon.
- **Tilt from horizontal** — 0° = camera pointing at horizon, 90° = camera pointing straight down.
- Internally, the app always converts to *tilt from nadir*.

### GSD Formula
```
GSD = (pixel_size_mm / focal_length_mm) × altitude_m / cos²(angle_from_nadir_rad)
```
- The `cos²` denominator accounts for:
  1. Increased slant range (`H / cos(α)`)
  2. Oblique foreshortening on the ground (second `1/cos(α)`)
- ⚠️ **Verify**: For highly oblique angles (>60°), this approximation becomes increasingly optimistic. True GSD should be verified with a full ray-cast model.

### Footprint
- **Across-track footprint** = far edge ground distance − near edge ground distance.
- **Along-track footprint** uses the slant range to the *image centre*:
  ```
  along_track = 2 × centre_slant × tan(half_fov_along)
  ```
  ⚠️ This is an approximation. For highly oblique cameras, the along-track footprint trapezoid is asymmetric. A full 4-corner projection would be more accurate.

### 4-Camera Oblique Arrangement
- In this app, all 4 cameras are modelled as **purely across-track oblique** (i.e., they tilt left/right only).
- ⚠️ Real 4-camera rigs (e.g. Leica RCD30 Oblique) may also have forward/rear angular components. If your rig has along-track tilt, the along-track footprint calculation needs extension.

### Line Spacing
```
line_spacing = combined_swath × (1 − sidelap_fraction)
```
- `combined_swath` = from the leftmost footprint near-edge to the rightmost far-edge.
- This ensures adjacent strip footprints overlap by the specified sidelap at the *footprint level*.

### Photo Spacing
```
photo_spacing = footprint_along × (1 − forward_overlap_fraction)
```
- Uses the **image centre** slant range to estimate along-track footprint.
- At the far edge, slant range is longer, so effective along-track footprint is larger — forward overlap there will be slightly higher than at the near edge.

### Reciprocal Flying
- For oblique cameras, flying the reciprocal (return) strip at a lateral offset fills gaps that one-direction strips leave.
- The app recommends reciprocal flying whenever the camera tilt is > 5°.

### Limitations / Future Work
- Terrain variation (DTM-based)
- Along-track oblique component (full 5-camera rigs)
- Lens distortion correction
- Wind drift / crab angle
- IMU/GPS lever-arm effects
    """)

# ---------------------------------------------------------------------------
# Export
# ---------------------------------------------------------------------------

st.subheader("📤 Export Results")
st.caption("Download the current survey summary and per-camera specs as Excel, Word, or PDF.")

import io
import datetime

def build_export_data():
    """Collect all summary values into plain dicts for export."""
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")

    summary = {
        "Export date": timestamp,
        "Camera name": camera_name,
        "Arrangement": arrangement,
        "Tilt from nadir (°)": round(tilt_from_nadir, 1),
        "Altitude AGL (m)": round(altitude_m, 1),
        "Aircraft speed (m/s)": round(speed_ms, 1),
        "Forward overlap (%)": forward_overlap_pct,
        "Sidelap (%)": sidelap_pct,
        "Sensor width (mm)": sensor_w_mm,
        "Sensor height (mm)": sensor_h_mm,
        "Image width (px)": img_w_px,
        "Image height (px)": img_h_px,
        "Focal length (mm)": focal_mm,
    }

    if mc:
        summary.update({
            "Combined swath (m)": round(mc.combined_swath_m, 1),
            "Line spacing (m)": round(mc.recommended_line_spacing_m, 1),
            "Photo spacing (m)": round(mc.recommended_photo_spacing_m, 1),
            "Exposure interval (s)": round(mc.photo_interval_s, 2),
            "Sidelap achieved (%)": round(mc.sidelap_achieved * 100, 1),
            "Fwd overlap near edge (%)": round(mc.forward_overlap_near * 100, 1),
            "Fwd overlap centre (%)": round(mc.forward_overlap_centre * 100, 1),
            "Fwd overlap far edge (%)": round(mc.forward_overlap_far * 100, 1),
            "Reciprocal flying recommended": "Yes" if mc.reciprocal_recommended else "No",
        })

    if cam_solutions:
        sol0 = cam_solutions[0]
        summary.update({
            "Near-edge GSD (cm/px)": round(sol0.near_gsd_m * 100, 2),
            "Centre GSD (cm/px)": round(sol0.centre_gsd_m * 100, 2),
            "Far-edge GSD (cm/px)": round(sol0.far_gsd_m * 100, 2),
            "Near angle from vertical (°)": round(sol0.near_angle_deg, 1),
            "Near angle from horizontal (°)": round(90.0 - sol0.near_angle_deg, 1),
            "Centre angle from vertical (°)": round(sol0.centre_angle_deg, 1),
            "Centre angle from horizontal (°)": round(90.0 - sol0.centre_angle_deg, 1),
        })

    per_camera = []
    for i, (lbl, sol) in enumerate(zip(labels, cam_solutions)):
        per_camera.append({
            "Camera": lbl,
            "Side": "Right" if tilts[i] >= 0 else "Left",
            "Near Edge (m)": round(sol.near_edge_m, 1),
            "Centre (m)": round(sol.centre_m, 1),
            "Far Edge (m)": round(sol.far_edge_m, 1),
            "Near Slant (m)": round(sol.near_slant_m, 1),
            "Centre Slant (m)": round(sol.centre_slant_m, 1),
            "Far Slant (m)": round(sol.far_slant_m, 1),
            "Near angle from vertical (°)": round(sol.near_angle_deg, 1),
            "Near angle from horizontal (°)": round(90.0 - sol.near_angle_deg, 1),
            "Centre angle from vertical (°)": round(sol.centre_angle_deg, 1),
            "Centre angle from horizontal (°)": round(90.0 - sol.centre_angle_deg, 1),
            "Near GSD (cm/px)": round(sol.near_gsd_m * 100, 2),
            "Centre GSD (cm/px)": round(sol.centre_gsd_m * 100, 2),
            "Far GSD (cm/px)": round(sol.far_gsd_m * 100, 2),
            "Footprint Across (m)": round(sol.footprint_across_m, 1),
            "Footprint Along (m)": round(sol.footprint_along_m, 1),
            "Pixel Size (µm)": round(sol.pixel_size_mm * 1000, 2),
            "Half FOV Across (°)": round(sol.half_fov_across_deg, 2),
        })

    return summary, per_camera


def make_excel(summary, per_camera):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()

    # --- Summary sheet ---
    ws1 = wb.active
    ws1.title = "Summary"
    hdr_font = Font(bold=True, color="FFFFFF", name="Arial", size=11)
    hdr_fill = PatternFill("solid", start_color="2E4057")
    key_font = Font(bold=True, name="Arial", size=10)
    val_font = Font(name="Arial", size=10)
    thin = Side(style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    ws1.append(["Oblique Survey Planner — Summary"])
    ws1["A1"].font = Font(bold=True, name="Arial", size=14)
    ws1.append([])

    ws1.append(["Parameter", "Value"])
    for cell in ws1[3]:
        cell.font = hdr_font
        cell.fill = hdr_fill
        cell.alignment = Alignment(horizontal="center")
        cell.border = border

    for row_idx, (k, v) in enumerate(summary.items(), start=4):
        ws1.append([k, v])
        ws1.cell(row=row_idx, column=1).font = key_font
        ws1.cell(row=row_idx, column=2).font = val_font
        for col in [1, 2]:
            ws1.cell(row=row_idx, column=col).border = border

    ws1.column_dimensions["A"].width = 36
    ws1.column_dimensions["B"].width = 28

    # --- Per-camera sheet ---
    ws2 = wb.create_sheet("Per-Camera Geometry")
    if per_camera:
        headers = list(per_camera[0].keys())
        ws2.append(headers)
        for cell in ws2[1]:
            cell.font = hdr_font
            cell.fill = hdr_fill
            cell.alignment = Alignment(horizontal="center")
            cell.border = border
        for row in per_camera:
            ws2.append(list(row.values()))
        for col_idx in range(1, len(headers) + 1):
            ws2.column_dimensions[get_column_letter(col_idx)].width = 22

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


def make_word(summary, per_camera):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Title
    title = doc.add_heading("Oblique Survey Planner — Export", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {summary.get('Export date', '')}")
    doc.add_paragraph("")

    # Summary table
    doc.add_heading("Survey Summary", level=1)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr_cells = tbl.rows[0].cells
    hdr_cells[0].text = "Parameter"
    hdr_cells[1].text = "Value"
    for cell in hdr_cells:
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "2E4057")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

    for k, v in summary.items():
        row_cells = tbl.add_row().cells
        row_cells[0].text = str(k)
        row_cells[1].text = str(v)

    tbl.columns[0].width = Inches(3.2)
    tbl.columns[1].width = Inches(2.8)

    doc.add_paragraph("")

    # Per-camera table
    if per_camera:
        doc.add_heading("Per-Camera Geometry", level=1)
        headers = list(per_camera[0].keys())
        tbl2 = doc.add_table(rows=1, cols=len(headers))
        tbl2.style = "Table Grid"
        hdr_cells2 = tbl2.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells2[i].text = h
            run = hdr_cells2[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tc = hdr_cells2[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "2E4057")
            shd.set(qn("w:val"), "clear")
            tcPr.append(shd)
        for row in per_camera:
            row_cells = tbl2.add_row().cells
            for i, v in enumerate(row.values()):
                row_cells[i].text = str(v)
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def make_pdf(summary, per_camera):
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=1.5*cm, rightMargin=1.5*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    hdr_style = ParagraphStyle("hdr", parent=styles["Heading1"],
                               fontSize=16, spaceAfter=6)
    sub_style = ParagraphStyle("sub", parent=styles["Heading2"],
                               fontSize=12, spaceAfter=4)
    body_style = styles["Normal"]

    story = []
    story.append(Paragraph("Oblique Survey Planner — Export", hdr_style))
    story.append(Paragraph(f"Generated: {summary.get('Export date', '')}", body_style))
    story.append(Spacer(1, 0.4*cm))

    # Summary table
    story.append(Paragraph("Survey Summary", sub_style))
    sum_data = [["Parameter", "Value"]] + [[k, str(v)] for k, v in summary.items()]
    sum_table = Table(sum_data, colWidths=[10*cm, 7*cm])
    sum_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2E4057")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F0F4F8")]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(sum_table)
    story.append(Spacer(1, 0.6*cm))

    # Per-camera table — landscape page
    if per_camera:
        story.append(PageBreak())
        story.append(Paragraph("Per-Camera Geometry", sub_style))
        headers = list(per_camera[0].keys())
        cam_data = [headers] + [list(r.values()) for r in per_camera]
        col_w = (A4[1] - 3*cm) / len(headers)  # landscape width minus margins
        cam_table = Table(cam_data, colWidths=[col_w] * len(headers), repeatRows=1)
        cam_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2E4057")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F0F4F8")]),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
            ("LEFTPADDING", (0, 0), (-1, -1), 3),
            ("RIGHTPADDING", (0, 0), (-1, -1), 3),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("WORDWRAP", (0, 0), (-1, -1), True),
        ]))
        story.append(cam_table)

    doc.build(story)
    buf.seek(0)
    return buf.read()


if cam_solutions:
    summary_data, per_camera_data = build_export_data()
    fname_base = scenario_name if scenario_name else "survey"

    ex1, ex2, ex3 = st.columns(3)

    with ex1:
        try:
            xlsx_bytes = make_excel(summary_data, per_camera_data)
            st.download_button(
                label="⬇️ Download Excel (.xlsx)",
                data=xlsx_bytes,
                file_name=f"{fname_base}_survey.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        except Exception as e:
            st.error(f"Excel export failed: {e}")

    with ex2:
        try:
            docx_bytes = make_word(summary_data, per_camera_data)
            st.download_button(
                label="⬇️ Download Word (.docx)",
                data=docx_bytes,
                file_name=f"{fname_base}_survey.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as e:
            st.error(f"Word export failed: {e}")

    with ex3:
        try:
            pdf_bytes = make_pdf(summary_data, per_camera_data)
            st.download_button(
                label="⬇️ Download PDF",
                data=pdf_bytes,
                file_name=f"{fname_base}_survey.pdf",
                mime="application/pdf",
            )
        except Exception as e:
            st.error(f"PDF export failed: {e}")

st.markdown("---")

# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------

st.markdown("---")
st.caption(
    "Oblique Survey Planner v1.0 | "
    "Flat terrain pinhole model | "
    "All formulae shown above — verify geometry assumptions before flight planning."
)
